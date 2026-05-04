"""Document -> per-page OCRPage[] dispatcher.

- Digital PDF (text layer present) -> pdfplumber, no OCR.
- Scanned PDF (no/empty text layer) -> rasterise pages -> OCREngine.
- Image (jpg/png/tiff) -> single-page OCREngine.
- DOCX -> python-docx text extraction.

Each page is upserted into mongo.ocr_pages keyed by (doc_id, page_number)
so re-running is idempotent.
"""
from __future__ import annotations

import io
from pathlib import Path

import pdfplumber
from PIL import Image

from app.db.mongo import get_mongo_db
from app.pipeline.ocr import OCRPage, OCRWord, get_ocr_engine

# Pages with fewer than this many text characters are treated as scanned.
DIGITAL_TEXT_MIN_CHARS = 30
# Render scanned pages at this DPI before OCR. 200 is a good sweet spot.
RASTER_DPI = 200


def _rasterise_pdf_page(pdf_path: str, page_number_zero_based: int) -> str:
    """Render one PDF page to PNG. Returns the PNG path."""
    # pdfplumber wraps pdfminer; for rasterising we use pypdfium2 which
    # ships transitively with pdfplumber.
    import pypdfium2 as pdfium

    pdf = pdfium.PdfDocument(pdf_path)
    page = pdf[page_number_zero_based]
    pil_image: Image.Image = page.render(scale=RASTER_DPI / 72).to_pil()
    out_path = Path(pdf_path).with_suffix(f".p{page_number_zero_based + 1}.png")
    pil_image.save(out_path, format="PNG")
    return str(out_path)


def _digital_page(plumber_page) -> OCRPage:
    text = plumber_page.extract_text() or ""
    return OCRPage(
        page_number=plumber_page.page_number,
        text=text,
        words=[],
        width=int(plumber_page.width or 0),
        height=int(plumber_page.height or 0),
        source="digital",
    )


def extract_pdf(pdf_path: str) -> list[OCRPage]:
    pages: list[OCRPage] = []
    with pdfplumber.open(pdf_path) as pdf:
        for idx, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            if len(text.strip()) >= DIGITAL_TEXT_MIN_CHARS:
                pages.append(_digital_page(page))
                continue
            # No (or near-empty) text layer -> rasterise + OCR.
            png_path = _rasterise_pdf_page(pdf_path, idx)
            ocr_page = get_ocr_engine().ocr_image(png_path, page_number=idx + 1)
            pages.append(ocr_page)
    return pages


def extract_image(image_path: str) -> list[OCRPage]:
    return [get_ocr_engine().ocr_image(image_path, page_number=1)]


def extract_docx(docx_path: str) -> list[OCRPage]:
    import docx as docx_lib  # python-docx

    doc = docx_lib.Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [
        OCRPage(
            page_number=1,
            text=text,
            words=[],
            source="digital",
        )
    ]


def extract_document(file_path: str) -> list[OCRPage]:
    suffix = Path(file_path).suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(file_path)
    if suffix in {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp"}:
        return extract_image(file_path)
    if suffix in {".docx", ".doc"}:
        return extract_docx(file_path)
    raise ValueError(f"unsupported_extension:{suffix}")


def persist_pages(doc_kind: str, doc_id: int, pages: list[OCRPage]) -> None:
    """Upsert each page into mongo.ocr_pages (idempotent on re-run)."""
    coll = get_mongo_db()["ocr_pages"]
    for p in pages:
        coll.update_one(
            {"doc_kind": doc_kind, "doc_id": doc_id, "page_number": p.page_number},
            {
                "$set": {
                    "doc_kind": doc_kind,
                    "doc_id": doc_id,
                    "page_number": p.page_number,
                    "text": p.text,
                    "source": p.source,
                    "width": p.width,
                    "height": p.height,
                    "avg_confidence": p.avg_confidence,
                    "words": [
                        {"t": w.text, "c": w.confidence, "b": list(w.bbox)}
                        for w in p.words
                    ],
                }
            },
            upsert=True,
        )
